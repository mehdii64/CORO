"""
Fills BACK UP M.docx template with FullReport data and returns BytesIO.
The template uses {{TOKEN}} placeholders that will be replaced in Task 10.
Until Task 10 is done (manual tokenization of the Word file), this builder
works but tokens won't be substituted visually — the text replacement is ready.
"""
from __future__ import annotations
import io, json
from pathlib import Path
from docx import Document
from services.phrase_builder import build_full_coro_text, build_full_atc_text

TEMPLATE_PATH = Path(__file__).parent.parent.parent / "templates" / "BACK UP M.docx"

def _s(value, default="") -> str:
    return str(value) if value is not None else default

def _format_fdrcx(clinical: dict) -> str:
    raw = clinical.get("fdrcx", "[]") or "[]"
    items = raw if isinstance(raw, list) else json.loads(raw)
    notes = clinical.get("fdrcx_notes", "") or ""
    result = ", ".join(items)
    if notes:
        result += f" — {notes}"
    return result

def _format_ett(clinical: dict) -> str:
    """Order: ETT description, then troubles de la cinétique, then FEVG."""
    fevg = clinical.get("ett_fevg")
    notes = (clinical.get("ett_notes") or "").strip()
    kinetics = (clinical.get("ett_kinetics") or "").strip()
    parts: list[str] = []
    if notes:
        parts.append(notes.rstrip(".") + ".")
    if kinetics:
        parts.append(kinetics.rstrip(".") + ".")
    if fevg:
        parts.append(f"FEVG {fevg}%.")
    return " ".join(parts) if parts else "VG de cinétique normale."

def _format_material(technique: dict) -> str:
    raw_mat = technique.get("material", "[]") or "[]"
    items = raw_mat if isinstance(raw_mat, list) else json.loads(raw_mat)
    # Items are joined with '\n'; the post-processor splits them into one
    # bulleted paragraph per item using the template's bullet formatting.
    return "\n".join(items)


def _material_items(technique: dict) -> list[str]:
    raw_mat = technique.get("material", "[]") or "[]"
    items = raw_mat if isinstance(raw_mat, list) else json.loads(raw_mat)
    return [str(x) for x in items]

import re

# Labels to strip from the template header (with or without trailing value).
# Match the label, optional spaces, optional ':', then anything up to a tab/newline/2+ spaces.
STRIP_LABEL_PATTERNS = [
    re.compile(r"Date de naissance\s*:?\s*"),
    re.compile(r"IPP\s*:?\s*[^\t\n]*"),
    re.compile(r"Taille\s*:?\s*[^\t\n]*"),
    re.compile(r"Poids\s*:?\s*[^\t\n]*"),
    re.compile(r"IMC\s*:?\s*[^\t\n]*"),
    re.compile(r"SC\s*\(m[²2]\)\s*:?\s*[^\t\n]*"),
]

# Tokens used to inject the missing "Artère Coronaire Droite:" header.
# We render it as bold, with the same look as "Artère Coronaire Gauche:".
ACD_HEADER_TEXT = "Artère Coronaire Droite : "


def _scrub_header_labels(para_text: str) -> str:
    out = para_text
    for pat in STRIP_LABEL_PATTERNS:
        out = pat.sub("", out)
    # collapse runs of tabs/spaces left behind
    out = re.sub(r"[ \t]{2,}", " ", out).strip()
    return out


def _rewrite_paragraph_text(para, new_text: str) -> None:
    """Replace the entire text of a paragraph while preserving its first run's style."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


def build_docx(full_report: dict) -> io.BytesIO:
    doc = Document(TEMPLATE_PATH)
    # Snapshot template paragraph styles before we mutate the document, so the
    # ATC page can clone the same look (blue header bands, body, bullets).
    templates = _capture_templates(doc)
    report   = full_report.get("report") or {}
    clinical = full_report.get("clinical") or {}
    technique= full_report.get("technique") or {}
    conclusion=full_report.get("conclusion") or {}

    texts = build_full_coro_text(full_report)

    doctor_names = full_report.get("doctor_names", [])

    dominance = full_report.get("dominance") or "droite"

    acd_text = texts.get("acd") or ""

    replacements = {
        "{{PATIENT_NAME}}":   _s(report.get("patient_name")),
        "{{DOB}}":            "",
        "{{SEX}}":            _s(report.get("sex")),
        "{{IPP}}":            "",
        "{{HEIGHT}}":         "",
        "{{WEIGHT}}":         "",
        "{{EXAM_DATE}}":      _s(report.get("exam_date")),
        "{{DATE}}":           _s(report.get("exam_date")),
        "{{INDICATION}}":     _s(report.get("indication")),
        "{{OPERATORS}}":      ", ".join(doctor_names),
        "{{AGE}}":            _s(clinical.get("age")),
        "{{FDRCX}}":          _format_fdrcx(clinical),
        "{{ATCD}}":           _s(clinical.get("atcd")),
        "{{HM}}":             _s(clinical.get("hm")),
        "{{ECG}}":            _s(clinical.get("ecg")),
        "{{ETT}}":            _format_ett(clinical),
        "{{ETT_KINETICS}}":   _s(clinical.get("ett_kinetics")),
        "{{FEVG}}":           f"{clinical.get('ett_fevg')}%" if clinical.get("ett_fevg") else "",
        "{{ROOM}}":           _s(technique.get("room")),
        "{{APPROACH}}":       _s(technique.get("approach")),
        "{{FRENCH}}":         _s(technique.get("french_size")),
        "{{PDC}}":            _s(technique.get("contrast_cc")),
        "{{DOSE}}":           _s(technique.get("dose_mgy")),
        # {{MATERIAL}} is handled separately (multi-paragraph bullet expansion).
        "{{CORO_TCG}}":       texts["tcg"],
        "{{CORO_IVA}}":       texts["iva"],
        "{{CORO_CX}}":        texts["cx"],
        "{{CORO_ACD}}":       acd_text,
        "{{DOMINANCE}}":      dominance,
        "{{DECISION}}":       _s(conclusion.get("decision")),
        "{{DECISION_NOTES}}": _s(conclusion.get("decision_notes")),
    }

    conclusion_lines = texts["conclusion_lines"]
    conclusion_idx = 0

    def replace_in_para(para):
        nonlocal conclusion_idx
        for run in para.runs:
            if "{{CONCLUSION}}" in run.text:
                if conclusion_idx < len(conclusion_lines):
                    run.text = run.text.replace("{{CONCLUSION}}", conclusion_lines[conclusion_idx])
                    conclusion_idx += 1
                else:
                    run.text = run.text.replace("{{CONCLUSION}}", "")
                continue
            for token, value in replacements.items():
                if token in run.text:
                    run.text = run.text.replace(token, value)

    for para in doc.paragraphs:
        replace_in_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    # Post-process: strip removed header labels and rename "Date Examen:" → "Date:".
    for para in doc.paragraphs:
        original = para.text
        new = original
        new = re.sub(r"Date\s+Examen\s*:", "Date :", new, flags=re.IGNORECASE)
        new = _scrub_header_labels(new)
        if new != original:
            _rewrite_paragraph_text(para, new)

    # Expand {{MATERIAL}} into one bulleted paragraph per item (cloning the
    # template paragraph's bullet formatting).
    _expand_material(doc, _material_items(technique))

    # Drop empty bulleted paragraphs left over from unused {{CONCLUSION}}/
    # {{MATERIAL}} placeholders so we don't render empty bullets.
    _remove_empty_bulleted_paragraphs(doc)

    # Inject "Artère Coronaire Droite :" header before the ACD paragraph,
    # if the template doesn't already include it.
    _inject_acd_header(doc, acd_text)

    # Append ATC section if applicable
    report_type = (report.get("type") or "coro").lower()
    if "atc" in report_type or full_report.get("interventions"):
        _append_atc_section(doc, full_report, replacements, templates)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _capture_templates(doc) -> dict:
    """Snapshot prototype paragraphs from the template so the ATC page can be
    rendered with the same visual style as the Coronarographie page (blue
    header bands, body lines, bulleted lists, centered title)."""
    from copy import deepcopy
    from docx.oxml.ns import qn
    out = {"title": None, "band": None, "body": None, "header_body": None, "bullet": None}
    for p in doc.paragraphs:
        txt = p.text.strip()
        pPr = p._p.find(qn("w:pPr"))
        has_shd = pPr is not None and pPr.find(qn("w:shd")) is not None
        has_num = pPr is not None and pPr.find(qn("w:numPr")) is not None
        if out["title"] is None and txt == "Coronarographie":
            out["title"] = deepcopy(p._p)
        if out["band"] is None and txt in ("TECHNIQUE", "CORONAROGRAPHIE", "CONCLUSION"):
            out["band"] = deepcopy(p._p)
        if out["body"] is None and txt.startswith("Salle"):
            out["body"] = deepcopy(p._p)
        if out["header_body"] is None and has_shd and not has_num and txt.startswith("Patient"):
            out["header_body"] = deepcopy(p._p)
        if out["bullet"] is None and has_num:
            out["bullet"] = deepcopy(p._p)
    return out


def _clone_with_text(template_elem, text: str, bold: bool | None = None):
    """Deepcopy a template <w:p>, replace text in the first run, drop other runs.
    If `bold` is False, strip `<w:b/>` from the run's rPr; if True, ensure it's set;
    if None, leave the template's bold setting as-is."""
    from copy import deepcopy
    from docx.oxml.ns import qn
    new_p = deepcopy(template_elem)
    runs = new_p.findall(qn("w:r"))
    if runs:
        first = runs[0]
        for tt in first.findall(qn("w:t")):
            first.remove(tt)
        tt = OxmlElement_safe("w:t")
        tt.text = text
        tt.set(qn("xml:space"), "preserve")
        first.append(tt)
        for r in runs[1:]:
            new_p.remove(r)
        if bold is not None:
            _set_run_bold(first, bold)
    else:
        r = OxmlElement_safe("w:r")
        tt = OxmlElement_safe("w:t")
        tt.text = text
        tt.set(qn("xml:space"), "preserve")
        r.append(tt)
        new_p.append(r)
        if bold is not None:
            _set_run_bold(r, bold)
    return new_p


def _clone_with_segments(template_elem, segments: list[tuple[str, bool]]):
    """Deepcopy a template <w:p>, then rebuild its content as a sequence of runs
    where each segment is (text, bold). The first template run's rPr is reused
    as the prototype run formatting (font/size/color), and bold is set per segment."""
    from copy import deepcopy
    from docx.oxml.ns import qn
    new_p = deepcopy(template_elem)
    runs = new_p.findall(qn("w:r"))
    proto_rPr = None
    if runs:
        proto_rPr = runs[0].find(qn("w:rPr"))
        for r in runs:
            new_p.remove(r)
    for text, bold in segments:
        r = OxmlElement_safe("w:r")
        if proto_rPr is not None:
            r.append(deepcopy(proto_rPr))
        tt = OxmlElement_safe("w:t")
        tt.text = text
        tt.set(qn("xml:space"), "preserve")
        r.append(tt)
        _set_run_bold(r, bold)
        new_p.append(r)
    return new_p


def _set_run_bold(run_elem, bold: bool) -> None:
    """Set or clear <w:b/> in a run's <w:rPr>."""
    from docx.oxml.ns import qn
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement_safe("w:rPr")
        run_elem.insert(0, rPr)
    for b in rPr.findall(qn("w:b")):
        rPr.remove(b)
    for b in rPr.findall(qn("w:bCs")):
        rPr.remove(b)
    if bold:
        b = OxmlElement_safe("w:b")
        rPr.append(b)


def _append_atc_section(doc, full_report: dict, base_replacements: dict, templates: dict) -> None:
    """Append the Angioplastie page using the same visual style as the
    Coronarographie page by cloning template paragraphs."""
    interventions = full_report.get("interventions", []) or []
    if not interventions:
        return

    atc_texts = build_full_atc_text(full_report)
    body = doc.element.body

    band_t   = templates.get("band")
    body_t   = templates.get("body")
    header_t = templates.get("header_body") or templates.get("body")
    title_t  = templates.get("title")
    bullet_t = templates.get("bullet")

    # Bold control per kind. Bands/titles keep the template's bold setting.
    # Body/bullet content goes through a single run with bold=False so the
    # narrative text isn't entirely bold. Mixed label/value lines use
    # `_append_segments` to render bold labels with non-bold values.
    bold_by_kind = {
        "band": None, "title": None,
        "body": False, "header_body": False,
        "bullet": False,
    }
    proto_by_kind = {
        "band": band_t, "body": body_t, "title": title_t,
        "bullet": bullet_t, "header_body": header_t,
    }

    def _append(text: str, kind: str) -> None:
        proto = proto_by_kind.get(kind)
        if proto is None:
            doc.add_paragraph(text)
            return
        body.append(_clone_with_text(proto, text, bold=bold_by_kind.get(kind)))

    def _append_segments(segments: list[tuple[str, bool]], kind: str) -> None:
        """Render a paragraph from `kind` template with mixed bold/non-bold runs."""
        proto = proto_by_kind.get(kind)
        if proto is None:
            doc.add_paragraph("".join(s for s, _ in segments))
            return
        body.append(_clone_with_segments(proto, segments))

    # Page break (cloned body para with break, to match page layout)
    doc.add_page_break()

    # Title
    _append("Angioplastie", "title")

    # Header info — same blue-shaded header block as the coronarography page
    pname = base_replacements.get("{{PATIENT_NAME}}", "")
    sex   = base_replacements.get("{{SEX}}", "")
    date  = base_replacements.get("{{EXAM_DATE}}", "")
    operators = base_replacements.get("{{OPERATORS}}", "")
    indication = base_replacements.get("{{INDICATION}}", "")
    _append_segments(
        [("Patient: ", True), (pname, False), ("	Sexe: ", True), (sex, False)],
        "header_body",
    )
    _append_segments([("Date : ", True), (date, False)], "header_body")
    if operators:
        _append_segments([("Opérateur: ", True), (operators, False)], "header_body")
    if indication:
        _append_segments([("Indications: ", True), (indication, False)], "header_body")

    # STATUT CLINIQUE band + bulleted clinical info (same as coro page)
    clinical = full_report.get("clinical") or {}
    _append("STATUT CLINIQUE", "band")
    _append_segments([("Âge: ", True), (_s(clinical.get("age")), False)], "bullet")
    _append_segments([("FDRCx : ", True), (_format_fdrcx(clinical), False)], "bullet")
    _append_segments([("ATCD : ", True), (_s(clinical.get("atcd")), False)], "bullet")
    _append_segments([("HM : ", True), (_s(clinical.get("hm")), False)], "bullet")
    _append_segments([("ECG  : ", True), (_s(clinical.get("ecg")), False)], "bullet")
    _append_segments([("ETT :  ", True), (_format_ett(clinical), False)], "bullet")

    # TECHNIQUE band + body
    _append("TECHNIQUE", "band")
    _append_segments(
        [("Salle: ", True), (f"{base_replacements.get('{{ROOM}}', '')}.", False)],
        "body",
    )
    approach = base_replacements.get("{{APPROACH}}", "")
    fr = base_replacements.get("{{FRENCH}}", "")
    if approach:
        _append_segments(
            [("Voie d’abord primaire:	", True), (f"{approach} - {fr} French.", False)],
            "body",
        )
    _append_segments([("Matériel :", True)], "body")
    for item in atc_texts.get("atc_material", []):
        _append(item, "bullet")

    # ANGIOPLASTIE band + bulleted narrative
    _append("ANGIOPLASTIE", "band")
    blocks = atc_texts.get("atc_narrative_blocks", [])
    for bi, block in enumerate(blocks):
        for line in block:
            _append(line, "bullet")
        if bi < len(blocks) - 1:
            _append("", "body")  # spacer between interventions

    # CONCLUSION band + bulleted conclusion
    _append("CONCLUSION", "band")
    for line in atc_texts.get("atc_conclusion_lines", []):
        _append(line, "bullet")

    if operators:
        _append("", "body")
        _append(operators, "body")


def Pt_safe(n):
    """Lazy import of docx.shared.Pt to avoid top-level coupling."""
    from docx.shared import Pt
    return Pt(n)


def _inject_acd_header(doc, acd_text: str) -> None:
    """Insert a bold 'Artère Coronaire Droite :' paragraph immediately before the
    paragraph containing the ACD text, mirroring the LCA layout. Idempotent: skips
    if the header already exists in the document.
    """
    if not acd_text:
        return
    # Already present?
    for p in doc.paragraphs:
        if "Artère Coronaire Droite" in p.text:
            return
    # Find the ACD paragraph by its text content.
    target = None
    snippet = acd_text[:40]
    for p in doc.paragraphs:
        if snippet and snippet in p.text:
            target = p
            break
    if target is None:
        return
    # Insert a new paragraph before `target`, mirroring the LCA header's paragraph
    # (style + run formatting) so size/underline match exactly.
    lca_para = next(
        (p for p in doc.paragraphs if "Artère Coronaire Gauche" in p.text and p.runs),
        None,
    )
    new_p = target.insert_paragraph_before("")
    new_p.style = lca_para.style if lca_para is not None else target.style
    run = new_p.add_run("Artère Coronaire Droite :")
    run.bold = True
    if lca_para is not None:
        ref = lca_para.runs[0]
        run.underline = ref.underline
        if ref.font.size is not None:
            from docx.shared import Emu
            run.font.size = Emu(int(ref.font.size))
        if ref.font.name:
            run.font.name = ref.font.name


def _find_material_paragraph(doc):
    """Return the paragraph still containing the {{MATERIAL}} token, if any."""
    for p in doc.paragraphs:
        if "{{MATERIAL}}" in p.text:
            return p
    return None


def _set_paragraph_text(para, text: str) -> None:
    """Set paragraph text on the first run, clearing any others (preserves rPr)."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for r in para.runs[1:]:
        r.text = ""


def _expand_material(doc, items: list[str]) -> None:
    """Replace the {{MATERIAL}} placeholder with one bulleted paragraph per item.
    Clones the template paragraph (preserving numPr/style) for each extra item.
    If items is empty, the placeholder paragraph is emptied (cleanup pass will
    drop it as an empty bullet)."""
    anchor = _find_material_paragraph(doc)
    if anchor is None:
        return
    if not items:
        _set_paragraph_text(anchor, "")
        return
    _set_paragraph_text(anchor, items[0])
    from copy import deepcopy
    from docx.oxml.ns import qn
    prev = anchor._p
    for item in items[1:]:
        new_p = deepcopy(anchor._p)
        # Keep first run (with its rPr), set its text, drop other runs.
        runs = new_p.findall(qn("w:r"))
        if runs:
            first = runs[0]
            for t in first.findall(qn("w:t")):
                first.remove(t)
            t = OxmlElement_safe("w:t")
            t.text = item
            t.set(qn("xml:space"), "preserve")
            first.append(t)
            for r in runs[1:]:
                new_p.remove(r)
        prev.addnext(new_p)
        prev = new_p


def OxmlElement_safe(tag: str):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def _remove_empty_bulleted_paragraphs(doc) -> None:
    """Remove paragraphs whose visible text is empty AND that carry a numPr
    (Word list bullet/number). These show up as stray empty bullets when
    pre-allocated placeholders aren't filled."""
    from docx.oxml.ns import qn
    for p in list(doc.paragraphs):
        if p.text.strip():
            continue
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        if pPr.find(qn("w:numPr")) is None:
            continue
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)
