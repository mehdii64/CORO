"""
Tokenise BACK UP M.docx en remplaçant les valeurs patient par des {{TOKENS}}.
Produit templates/BACK UP M.docx (écrase l'original après backup).
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import shutil, copy

TEMPLATE = Path(__file__).parent.parent / "templates" / "BACK UP M.docx"
BACKUP   = Path(__file__).parent.parent / "templates" / "BACK UP M.original.docx"

def _keep_bold_replace_value(para, token: str):
    """
    Dans un paragraphe label:valeur, garde les runs gras (label)
    et remplace tous les runs non-gras par un seul run contenant le token.
    """
    if not para.runs:
        return
    # Trouve le premier run non-gras = début de la valeur
    value_start = None
    for i, r in enumerate(para.runs):
        if not r.bold:
            value_start = i
            break
    if value_start is None:
        # Tout est gras : remplace le dernier run
        para.runs[-1].text = " " + token
        return
    # Copie le format du premier run valeur pour le conserver
    first_value_run = para.runs[value_start]
    first_value_run.text = " " + token
    # Supprime tous les runs valeur suivants
    for r in para.runs[value_start + 1:]:
        r.text = ""


def _replace_all_runs(para, token: str):
    """Remplace tout le contenu du paragraphe par le token (1 seul run)."""
    if not para.runs:
        return
    para.runs[0].text = token
    for r in para.runs[1:]:
        r.text = ""


def _replace_run_text(run, old: str, new: str):
    if old in run.text:
        run.text = run.text.replace(old, new)


def tokenize():
    # Backup
    if not BACKUP.exists():
        shutil.copy(TEMPLATE, BACKUP)
        print(f"Backup créé : {BACKUP.name}")

    doc = Document(TEMPLATE)
    paras = doc.paragraphs

    for i, p in enumerate(paras):
        txt = p.text

        # ── En-tête patient ─────────────────────────────────────────────
        if txt.startswith("Patient:") and "EL OBIDI" in txt:
            # Run 2 = nom patient
            for r in p.runs:
                if "EL OBIDI MOHAMMED" in r.text:
                    r.text = r.text.replace("EL OBIDI MOHAMMED", "{{PATIENT_NAME}}")
                if "Masculin" in r.text:
                    r.text = r.text.replace("Masculin", "{{SEX}}")

        elif txt.startswith("IPP:") and "2126" in txt:
            for r in p.runs:
                if "2126" in r.text:
                    r.text = r.text.replace("2126", "{{IPP}}")

        elif txt.startswith("Date Examen:") and "2026" in txt:
            for r in p.runs:
                if "14/04/2026" in r.text:
                    r.text = r.text.replace("14/04/2026", "{{EXAM_DATE}}")

        elif txt.startswith("Op") and "Dr NAJLAA" in txt:
            # Runs après "Opérateur:" → operators
            for r in p.runs:
                if r.bold is None or not r.bold:
                    if "Dr" in r.text or "BELHARTY" in r.text or "HAJJAJ" in r.text:
                        r.text = "{{OPERATORS}}"
                    elif r.text.strip() and r.text.strip() not in (":",""):
                        r.text = ""

        elif txt.startswith("Indications:") or txt.startswith("Indications :"):
            # Value run may be bold too — target run[1] (after label run[0])
            non_bold = [r for r in p.runs if not r.bold and r.text.strip() and r.text.strip() != ":"]
            if non_bold:
                non_bold[0].text = " {{INDICATION}}"
                for r in non_bold[1:]:
                    r.text = ""
            elif len(p.runs) >= 2:
                p.runs[1].text = " {{INDICATION}}"
                for r in p.runs[2:]:
                    r.text = ""

        # ── Statut clinique ──────────────────────────────────────────────
        elif "Âge" in txt or "ge:" in txt.lower() and "73" in txt:
            # Âge: 73
            for r in p.runs:
                if not r.bold and ("73" in r.text or "\xa073" in r.text):
                    r.text = "\xa0{{AGE}}"
                    break

        elif txt.startswith("FDRCx") or "FDRCx" in txt[:8]:
            # Fusionne les runs valeur (non-gras) en un seul token
            value_runs = [r for r in p.runs if not r.bold]
            if value_runs:
                # Le premier run non-gras contient ": valeur..." → remplace
                value_runs[0].text = ":\xa0{{FDRCX}}"
                for r in value_runs[1:]:
                    r.text = ""

        elif txt.startswith("ATCD") and ":" in txt:
            value_runs = [r for r in p.runs if not r.bold]
            if value_runs:
                value_runs[0].text = ": {{ATCD}}"
                for r in value_runs[1:]:
                    r.text = ""

        elif txt.startswith("HM") and ":" in txt:
            value_runs = [r for r in p.runs if not r.bold]
            if value_runs:
                value_runs[0].text = ":\xa0{{HM}}"
                for r in value_runs[1:]:
                    r.text = ""

        elif txt.startswith("ECG") and ":" in txt and "RRS" in txt:
            value_runs = [r for r in p.runs if not r.bold]
            if value_runs:
                value_runs[0].text = "\xa0:\xa0{{ECG}}"
                for r in value_runs[1:]:
                    r.text = ""

        elif txt.startswith("ETT") and ":" in txt:
            value_runs = [r for r in p.runs if not r.bold]
            if value_runs:
                value_runs[0].text = " {{ETT}}"
                for r in value_runs[1:]:
                    r.text = ""

        # ── Technique ───────────────────────────────────────────────────
        elif txt.startswith("Salle:") and "innova" in txt.lower():
            for r in p.runs:
                if not r.bold and ("innova" in r.text.lower() or "GE" in r.text):
                    r.text = "{{ROOM}}. "
                    break
                elif not r.bold and r.text.strip() in ("", "\t"):
                    pass

        elif txt.startswith("Voie d") and "French" in txt:
            for r in p.runs:
                if "f" in r.text.lower() and "morale" in r.text.lower():
                    r.text = "{{APPROACH}} - {{FRENCH}} French. \xa0"
                elif "120" in r.text:
                    r.text = "{{PDC}} cc \xa0"
                elif "285" in r.text or "mGy" in r.text:
                    r.text = r.text.replace("285", "{{DOSE}}")
                elif "120 cc" in r.text:
                    r.text = r.text.replace("120 cc", "{{PDC}} cc")

        elif txt.startswith("Mat") and "riel" in txt and ":" in txt:
            pass  # Garde le label "Matériel:"

        elif txt.strip() in ("0.035 (ABBOTT LABORATORIES).",
                              "FR 4 (Boston Scientific).",
                              "IMPULSE 5F JL 3.5 (Boston Scientific).",
                              "Introducer F\xe9moral 6F (Terumo).",
                              "Introducer F�moral 6F (Terumo).") \
             or ("LABORATORIES" in txt or "Boston Scientific" in txt
                 or "Terumo" in txt):
            # Première ligne matériel → token, les suivantes vidées
            if i > 0 and "LABORATORIES" in txt:
                _replace_all_runs(p, "{{MATERIAL}}")
            elif i > 0 and any(x in txt for x in ["Boston", "Terumo", "Medtronic"]):
                _replace_all_runs(p, "")

        # ── Coronarographie ──────────────────────────────────────────────
        elif txt.startswith("Dominance:"):
            for r in p.runs:
                if not r.bold and "droite" in r.text.lower():
                    r.text = r.text.replace("Dominance droite", "{{DOMINANCE}}")
                    r.text = r.text.replace("Dominance gauche", "{{DOMINANCE}}")
                    r.text = r.text.replace("Dominance \xe9quilibr\xe9e", "{{DOMINANCE}}")

        elif txt.startswith("Le tronc commun se connecte"):
            _replace_all_runs(p, "{{CORO_TCG}}")

        elif txt.startswith("L'art") and "interventriculaire" in txt:
            _replace_all_runs(p, "{{CORO_IVA}}")

        elif txt.startswith("L'art") and "circonf" in txt:
            _replace_all_runs(p, "{{CORO_CX}}")

        elif "coronaire droite se connecte" in txt:
            _replace_all_runs(p, "{{CORO_ACD}}")

        # ── Conclusion ───────────────────────────────────────────────────
        elif txt.startswith("Coro") and ":" in txt:
            # Garde le label "Coro :", vide les runs valeur
            for r in p.runs:
                if not r.bold:
                    r.text = ""

        elif any(txt.startswith(x) for x in ["L\xe9sion", "L’sion",
                  "Lésion", "Lésion", "Lésions"]):
            _replace_all_runs(p, "{{CONCLUSION}}")

        elif txt.startswith("D\xe9cision") or txt.startswith("Décision"):
            # All runs may be bold; run[0] contains label+value mixed
            if p.runs:
                r0 = p.runs[0]
                raw = r0.text
                colon_pos = raw.find(":")
                if colon_pos != -1:
                    r0.text = raw[:colon_pos + 1] + " {{DECISION}}"
                else:
                    r0.text = "D\xe9cision\xa0: {{DECISION}}"
                for r in p.runs[1:]:
                    r.text = ""

        elif txt.strip().startswith("Dr ") and i > 40:
            # Ligne opérateurs en bas du CR
            _replace_all_runs(p, "{{OPERATORS}}")

    # Sauvegarde
    doc.save(TEMPLATE)
    print(f"[OK] Template tokenise : {TEMPLATE}")
    print()

    # Vérification
    doc2 = Document(TEMPLATE)
    tokens_found = []
    all_text = " ".join(p.text for p in doc2.paragraphs)
    for token in ["{{PATIENT_NAME}}", "{{EXAM_DATE}}", "{{IPP}}", "{{SEX}}",
                  "{{OPERATORS}}", "{{INDICATION}}", "{{AGE}}", "{{FDRCX}}",
                  "{{ATCD}}", "{{HM}}", "{{ECG}}", "{{ETT}}", "{{ROOM}}",
                  "{{APPROACH}}", "{{FRENCH}}", "{{PDC}}", "{{DOSE}}",
                  "{{MATERIAL}}", "{{CORO_TCG}}", "{{CORO_IVA}}",
                  "{{CORO_CX}}", "{{CORO_ACD}}", "{{CONCLUSION}}", "{{DECISION}}"]:
        if token in all_text:
            tokens_found.append(token)
        else:
            print(f"  [!] Token manquant : {token}")
    print(f"Tokens présents : {len(tokens_found)}/24")

if __name__ == "__main__":
    tokenize()
